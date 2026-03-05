import os
from typing import List, Optional
from langchain_core.documents import Document as LCDocument
from docx import Document as DocxDocument
from pypdf import PdfReader

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.md', '.markdown'}


def _load_pdf(file_path: str) -> Optional[str]:
    """从 PDF 文件中提取文本内容"""
    try:
        reader = PdfReader(file_path)
        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
        if not pages:
            return None
        return "\n".join(pages)
    except Exception as e:
        print(f"Failed to read PDF {file_path}: {e}")
        return None


def _load_docx(file_path: str) -> Optional[str]:
    """从 docx 文件中提取文本内容"""
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return None
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Failed to read DOCX {file_path}: {e}")
        return None


def _load_markdown(file_path: str) -> Optional[str]:
    """从 Markdown 文件中读取文本内容"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        return content if content else None
    except Exception as e:
        print(f"Failed to read Markdown {file_path}: {e}")
        return None


def _get_file_extension(file_path: str) -> str:
    """获取文件扩展名(小写)"""
    return os.path.splitext(file_path)[1].lower()


def _load_single_document(file_path: str) -> Optional[LCDocument]:
    """根据文件扩展名加载单个文档，支持 pdf、doc/docx、markdown"""
    extension = _get_file_extension(file_path)
    if extension not in SUPPORTED_EXTENSIONS:
        return None

    content = None
    if extension == ".pdf":
        content = _load_pdf(file_path)
    elif extension in (".doc", ".docx"):
        content = _load_docx(file_path)
    elif extension in (".md", ".markdown"):
        content = _load_markdown(file_path)

    if not content:
        return None

    file_name = os.path.basename(file_path)
    metadata = {
        "source": file_path,
        "file_name": file_name,
        "file_type": extension.lstrip("."),
        "total_chars": len(content),
    }

    return LCDocument(page_content=content, metadata=metadata)


def load_document(load_path: str) -> List[LCDocument]:
    """
    加载指定路径的文档，支持 pdf、doc/docx、markdown 格式。
    支持单个文件路径或文件夹路径(递归加载所有支持格式的文件)。
    """
    if not os.path.exists(load_path):
        print(f"Path does not exist: {load_path}")
        return []

    docs = []
    if os.path.isdir(load_path):
        for root, _, files in os.walk(load_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                if _get_file_extension(file_path) in SUPPORTED_EXTENSIONS:
                    doc = _load_single_document(file_path)
                    if doc:
                        docs.append(doc)
    else:
        doc = _load_single_document(load_path)
        if doc:
            docs.append(doc)

    print(f"Loaded {len(docs)} documents from {load_path}")
    return docs
