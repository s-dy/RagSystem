import os
import json
from typing import List, Optional
from langchain_core.documents import Document as LCDocument
from docx import Document as DocxDocument


def _load_document_with_metadata(file_path: str, metadata_dict: dict) -> Optional[LCDocument]:
    """
    从 .docx 加载内容，并从 metadata_dict 中补充元信息
    """
    if not file_path.endswith('.docx'):
        return None

    file_id = os.path.basename(file_path).replace('.docx', '')

    meta = metadata_dict.get(file_id, {})

    # 读取 Word 内容
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return None

        content = "\n".join(paragraphs)
    except Exception as e:
        print(f"❌ Failed to read {file_path}: {e}")
        return None

    # 构建元数据
    metadata = {
        "file_id": file_id,
        **meta
    }

    return LCDocument(page_content=content, metadata=metadata)

def _load_metadata_dict(data_dir) -> dict:
    """加载元信息"""
    metadata_file = os.path.join(data_dir, "metadata.json")

    if os.path.exists(metadata_file):
        try:
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata_dict = json.load(f)
        except Exception as e:
            print(f"⚠️ Failed to load metadata.json: {e}")
            metadata_dict = {}
    else:
        metadata_dict = {}
        print("⚠️ No metadata.json found")
    return metadata_dict

def load_document(load_path: str) -> List[LCDocument]:
    """
    加载 load_path 的 .docx 文档，并关联 metadata.json
    支持 单个文件具体路径（加载单个文档） 和 文件夹路径（加载所有文档）
    """
    if not os.path.exists(load_path):
        return []
    # 加载元信息
    if os.path.isdir(load_path):
        data_dir = load_path
    else:
        data_dir = os.path.dirname(load_path)
    metadata_dict = _load_metadata_dict(data_dir)
    # 加载文档
    docs = []
    if os.path.isdir(load_path):
        for file in os.listdir(data_dir):
            if file.endswith('.docx'):
                path = os.path.join(data_dir, file)
                doc = _load_document_with_metadata(path, metadata_dict)
                if doc:
                    docs.append(doc)
    else:
        doc = _load_document_with_metadata(load_path, metadata_dict)
        if doc:
            docs.append(doc)

    print(f"✅ Loaded {len(docs)} documents from {data_dir}")
    return docs


# 使用示例
if __name__ == '__main__':
    docs = load_document("/Users/sdy/hybridRag/src/crawler/data/0a58368b-332d-4982-85f4-330853baf00e.docx")
    # docs = load_document("/Users/sdy/hybridRag/src/crawler/data")
    if docs:
        doc = docs[0]
        print(doc.metadata)
